import io
import os
from typing import Union

try:
    from docx import Document
except Exception:
    Document = None


def extract_text_from_docx(docx_bytes_or_path: Union[bytes, str]) -> str:
    """Extract text from a .docx file given as bytes or file path.

    Requires the `python-docx` package (importable as `docx`).
    """
    if Document is None:
        raise RuntimeError("python docx is required to parse .docx files; install with `pip install python-docx`")

    if isinstance(docx_bytes_or_path, (bytes, bytearray)):
        fh = io.BytesIO(docx_bytes_or_path)
        doc = Document(fh)
    else:
        doc = Document(docx_bytes_or_path)

    paragraphs = [p.text for p in doc.paragraphs]
    return "\n".join(paragraphs)


def extract_text_from_txt(txt_bytes_or_path: Union[bytes, str]) -> str:
    """Extract text from a plain text or markdown file given as bytes or file path."""
    if isinstance(txt_bytes_or_path, (bytes, bytearray)):
        try:
            return txt_bytes_or_path.decode("utf-8")
        except Exception:
            return txt_bytes_or_path.decode("latin-1", errors="ignore")
    else:
        with open(txt_bytes_or_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()


def extract_text(file_or_bytes: Union[bytes, str], filename: str = None) -> str:
    """Unified extractor for DOCX and text files. Delegates PDFs to `pdf_parser` when detected.

    - If `file_or_bytes` is a path string, the function picks a parser based on the file extension.
    - If `file_or_bytes` is bytes, provide `filename` to help with detection (optional).
    """
    # Determine extension
    ext = None
    if filename:
        _, ext = os.path.splitext(filename.lower())
    elif isinstance(file_or_bytes, str):
        _, ext = os.path.splitext(file_or_bytes.lower())

    if ext == ".docx":
        return extract_text_from_docx(file_or_bytes)
    if ext in (".txt", ".md"):
        return extract_text_from_txt(file_or_bytes)
    if ext == ".pdf":
        # avoid circular import at module load time
        from .pdf_parser import extract_text_from_pdf

        return extract_text_from_pdf(file_or_bytes)

    # If no extension or unknown: try DOCX first, then text, then PDF
    if isinstance(file_or_bytes, (bytes, bytearray)):
        try:
            return extract_text_from_docx(file_or_bytes)
        except Exception:
            pass
        try:
            return extract_text_from_txt(file_or_bytes)
        except Exception:
            pass
        # fallback to PDF
        from .pdf_parser import extract_text_from_pdf

        return extract_text_from_pdf(file_or_bytes)

    # file_or_bytes is a path string but extension unknown -> try in order
    try:
        return extract_text_from_docx(file_or_bytes)
    except Exception:
        pass
    try:
        return extract_text_from_txt(file_or_bytes)
    except Exception:
        pass
    from .pdf_parser import extract_text_from_pdf

    return extract_text_from_pdf(file_or_bytes)
